import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_thesis_word_doc():
    doc = Document()
    
    # Set default font to Times New Roman, 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title Page
    doc.add_heading('ScoutAnt: An Intelligent Machine Learning System for Valorant Esports Analytics, Match Prediction, and NLP-Driven Insights', 0)
    
    subtitle = doc.add_paragraph('A Thesis Submitted in Partial Fulfillment of the Requirements for the Degree')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n\n\n')
    author = doc.add_paragraph('By\nDhruv\n\n2026')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # Read Markdown File
    md_path = 'c:\\Users\\dhruv\\ScoutAnt\\ScoutAnt_Thesis.md'
    if os.path.exists(md_path):
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        in_code_block = False
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('```'):
                in_code_block = not in_code_block
                continue
                
            if in_code_block:
                p = doc.add_paragraph(line)
                p.style = doc.styles['No Spacing']
                p.runs[0].font.name = 'Courier New'
                p.runs[0].font.size = Pt(10)
                continue
                
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. ') or line.startswith('2. '):
                doc.add_paragraph(line[3:], style='List Number')
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.line_spacing = 2.0  # Double spacing

    doc.add_page_break()
    
    # Appendix D: Source Code
    doc.add_heading('Appendix D: Complete System Source Code', level=1)
    p = doc.add_paragraph('The following appendix contains the complete source code for the ScoutAnt system, including the data scraper, machine learning pipeline, analytics API, and legacy tools.')
    p.paragraph_format.line_spacing = 2.0
    
    base_dir = 'c:\\Users\\dhruv\\ScoutAnt'
    for root, dirs, files in os.walk(base_dir):
        if '.git' in root or '.venv' in root or '__pycache__' in root:
            continue
        for file in files:
            if file.endswith('.py') or file.endswith('.json') or file.endswith('.css') or file.endswith('.ts') or file.endswith('.tsx'):
                # Limit JSON files to avoid totally breaking word
                if file.endswith('.json') and 'package-lock' in file:
                    continue
                    
                file_path = os.path.join(root, file)
                
                try:
                    with open(file_path, 'r', encoding='utf-8') as code_f:
                        content = code_f.read()
                        
                    # Truncate overly massive json
                    if file.endswith('.json') and len(content) > 100000:
                        content = content[:100000] + '\n... [TRUNCATED] ...'
                        
                    doc.add_heading(f'File: {os.path.relpath(file_path, base_dir)}', level=3)
                    cp = doc.add_paragraph(content)
                    cp.style = doc.styles['No Spacing']
                    
                    # Try to set font, but handle large blocks carefully
                    for run in cp.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(8)
                        
                except Exception as e:
                    pass

    # Save document
    output_path = 'c:\\Users\\dhruv\\ScoutAnt\\ScoutAnt_Thesis.docx'
    doc.save(output_path)
    print(f"Thesis successfully saved to {output_path}")

if __name__ == '__main__':
    create_thesis_word_doc()
